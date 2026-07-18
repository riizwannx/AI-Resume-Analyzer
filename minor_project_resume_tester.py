import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API key kaha hai bhai")

client=Groq(api_key=my_api_key)
model = "openai/gpt-oss-120b"


job_description="""
Description
The Full Stack Developer – Corporate Risk Reporting role focuses on the risk visualization layer of the Corporate Risk Reporting (CRR) platform. The role is responsible for building application capabilities that enable fast, flexible, and intuitive consumption of enterprise risk data for internal risk management users.
Within the CRR architecture, core risk logic is implemented in PostgreSQL, with curated risk results published into OpenSearch to support high-performance reporting, fuzzy filtering, and interactive exploration. This role remains hands-on in building solutions using .NET-based backend services and React-based frontend applications, while also leading and coaching junior full stack developers to deliver cohesive, high-quality visualization capabilities.
The position plays a key role in translating complex risk datasets into scalable and user-friendly reporting experiences, while setting technical direction and standards for the CRR visualization layer.
Key Accountabilities
RISK VISUALIZATION DELIVERY:
Designs and implements user-facing risk reporting capabilities that enable efficient visualization and exploration of enterprise risk data, including positions, P&L, FX exposures, and entering adjustments.

FRONTEND EXPERIENCE DELIVERY:
Development of high-performance and intuitive user interfaces using React and AgGrid, optimized for large datasets and advanced interaction patterns such as drilldowns, pivots, filtering, and navigation.

BACKEND SERVICES AND APIs:
Develops .NET-based backend services that mediate access to OpenSearch, support aggregation and query orchestration, and enable consistent data consumption patterns.

OPENSEARCH CONSUMPTION:
Implements effective OpenSearch query, indexing and consumption patterns to support fast retrieval, fuzzy filtering, and flexible exploration of risk data published from PostgreSQL.

ENGINEERING QUALITY:
Enforces engineering standards for the visualization layer, including code quality, testing practices, and CI/CD adoption, ensuring maintainable and reliable solutions.

PERFORMANCE AND USABILITY:
Owns performance and usability outcomes for the visualization layer, actively identifying and addressing bottlenecks across frontend, backend, and OpenSearch interactions.

COLLABORATION:
Works closely with the CRR Development Lead and Senior Business Analyst to align visualization capabilities with platform architecture and business risk requirements.

CONTINUOUS IMPROVEMENT:
Continuously evaluates and improves visualization patterns, technical approaches, and team practices to enhance the effectiveness and scalability of CRR risk reporting.

Qualifications
Minimum requirement of 3 years of relevant work experience, typically reflecting 5+ years in delivering full-stack applications in data-intensive or analytics-driven environments.

Demonstrated experience delivery of user-facing reporting or analytics solutions, with proven ability to guide and mentor other engineers.

Required Skills:

Mandatory:

Strong experience developing backend services using .NET
Strong experience building frontend applications using React
Hands-on experience with AgGrid for large, data-intensive reporting views
Exposure to cloud-native deployment models and CI/CD pipelines
Good to Have:

Experience working with risk, finance, or trading reporting platforms
Experience consuming and querying OpenSearch for fast, flexible reporting and fuzzy filtering
Familiarity with PostgreSQL-backed data solutions
Experience coaching or mentoring engineers in a delivery-focused environment
"""
class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f"""
Analyze the following job description:

{job_description}
"""
message_system={
    "role" : "system",
    "content" : system_prompt
}
message_user={
    "role" : "user",
    "content" : user_prompt
}
response_format={
    "type" : "json_object"
}


messages=[message_system, message_user]

response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)


answer=response.choices[0].message.content

raw_json=answer


import json
job_data=json.loads(raw_json)

job = JobD(**job_data)

print("minimum job experience requirements for this job->",job.minimum_experience)
print("minimum education requirements for this job ->",job.education_requirements)



#parse real
class MatchResult(BaseModel):
    score: float
    details: dict
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


resume_schema = Resume.model_json_schema()
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)
def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume


from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None


resume_folder = Path("resumes")
all_results=[]
for file_path in resume_folder.iterdir():
   
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume=parse_resume(resume_text) # llm call1
    time.sleep(5)
    result = final_score(job, parsed_resume) #llm caLL2
    
    time.sleep(5)
    print("Score:", result.score)
    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    })
all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)
top_2 = all_results[:2]
worst_2 = all_results[-2:]


print("\n","TOP 2 CANDIDATES AMONG ALL RESUME:")
for candidate in top_2:
    print(
        "\n",
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print("\n",candidate["details"])


print("\n","LOWEST 2 CANDIDATES AMONG ALL RESUME:")
for candidate in worst_2:
    
    print(
        "\n",
        candidate["name"],
        "-",
        candidate["score"],
        "%","\n"
    )
    
    print(candidate["details"])